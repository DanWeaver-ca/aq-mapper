"""Generate the Word guide (Home_Base_Guide.docx) for TAs / colleagues.

The printable twin of GETTING_STARTED.md — keep the two in sync.
Run:  python3 make_guide.py
Edit the content here and re-run to regenerate the document.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "Home_Base_Guide.docx")

doc = Document()

# Base font
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def shade(paragraph, fill="F3F3F3"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def code(text):
    """A shaded monospace command block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    shade(p)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def step(text):
    return doc.add_paragraph(text, style="List Number")


def note(text):
    p = doc.add_paragraph()
    r = p.add_run("Tip:  ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x77, 0x66)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(8)
    return p


# --- Title --------------------------------------------------------------
doc.add_heading("AQ Mapper — Home Base Guide", level=0)
sub = doc.add_paragraph()
sr = sub.add_run("How to set up and run the classroom air-quality dashboard. "
                 "Written for a TA or colleague — no programming needed. "
                 "Allow ~10 minutes for the one-time setup.")
sr.italic = True
sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()

# --- What this tool does ------------------------------------------------
doc.add_heading("1. What Home Base does", level=1)
doc.add_paragraph(
    "During the lab, student groups record air-quality readings on their "
    "phones with the AQ Mapper app and send you their data as CSV files. "
    "Home Base merges all of those files and builds one interactive "
    "dashboard (classroom_dashboard.html) you project for the debrief. "
    "It has four tabs:")
bullet(" every reading on a campus map. Filter to one group, outdoor or "
       "indoor, or everyone; colour the points by PM2.5, CO₂, temperature "
       "and more — the colours match the phone app exactly.", "Map:")
bullet(" a smooth “estimated field” between the points that fades "
       "where nobody sampled — with a smoothing-radius selector for the "
       "how-do-global-datasets-do-it discussion.", "Interpolated:")
bullet(" a PM2.5 density view of pollution hotspots.", "Heatmap:")
bullet(" indoor-vs-outdoor averages and a per-group table.", "Stats:")
doc.add_paragraph(
    "Above the tabs, a header shows the class totals and a group checklist: "
    "green chips for groups whose data is in, grey chips for groups still "
    "missing — so you can see who is outstanding while the class is running. "
    "The dashboard is a single ordinary HTML file: project it, and email it "
    "to students afterward.")

# --- What you'll need ---------------------------------------------------
doc.add_heading("2. What you’ll need", level=1)
bullet(" Windows or Mac, either is fine.", "A laptop:")
bullet(" the folder containing home_base.py and the other files (the same "
       "folder this guide came in).", "This tool:")
bullet(" a free program that runs the tool (one-time install, below).",
       "Python 3:")
bullet(" classroom wifi. Only the map’s background needs internet; "
       "everything else works offline.", "Internet:")
doc.add_paragraph(
    "You do not need real data to set up: a fake class of five groups is "
    "bundled, so you can rehearse everything beforehand.")

# --- One-time setup -----------------------------------------------------
doc.add_heading("3. One-time setup (once per computer)", level=1)

doc.add_heading("3.1  Install Python", level=2)
doc.add_paragraph("Windows:")
bullet(" Open the Microsoft Store, search for “Python 3”, and click "
       "Get/Install. (Or download from python.org and, on the first install "
       "screen, tick “Add Python to PATH”.)")
doc.add_paragraph("Mac:")
bullet(" Python 3 is usually already installed. If not, download it from "
       "python.org and run the installer.")
doc.add_paragraph("Already have Anaconda?")
bullet(" Anaconda is Python 3 — skip the install. But it keeps itself out "
       "of the normal terminal: open “Anaconda Prompt” from the Start "
       "menu and use that window for every command in this guide. (No "
       "Anaconda Prompt in the Start menu? Use the full-path row in the "
       "troubleshooting table.)")
doc.add_paragraph("To confirm it worked, open a terminal "
                  "(Windows: “Command Prompt”; Mac: "
                  "“Terminal”) and type:")
code("python --version")
note("If Windows says ‘python is not recognized’, try ‘py "
     "--version’. On Mac, use ‘python3 --version’. Use that "
     "same word (python / py / python3) everywhere below.")

doc.add_heading("3.2  Install the libraries the tool uses", level=2)
doc.add_paragraph(
    "In the terminal, move into this tool’s folder and install its "
    "requirements. To move into the folder, type ‘cd ’ (with a "
    "space) and then drag the folder onto the terminal window and press "
    "Enter. Then run:")
code("pip install -r requirements.txt")
note("If ‘pip’ isn’t found, try ‘pip3 install -r "
     "requirements.txt’ or ‘python -m pip install -r "
     "requirements.txt’. You only ever do this once per computer.")

doc.add_heading("3.3  Rehearse with the bundled sample data", level=2)
doc.add_paragraph(
    "Double-click the launcher — run_windows.bat (Windows) or "
    "run_mac.command (Mac; the first time only, right-click it and choose "
    "Open, then Open again in the dialog). With no real CSVs present it "
    "automatically uses the bundled fake class and opens the dashboard in "
    "your browser. (The Windows launcher finds Python on its own: PATH, "
    "then ‘py’, then a standard Anaconda install.) If you see the "
    "header, a row of green group chips, and the four tabs, you are ready "
    "for lab day.")

# --- Each lab -----------------------------------------------------------
doc.add_heading("4. Each lab: getting the data in", level=1)
step("Have each group send their data: in the AQ Mapper app they open the "
     "Data screen and tap “Send to instructor”, then AirDrop or "
     "email you the file (named like aq_UTSC-AQMS-07_20260716_143022.csv).")
step("Collect all the groups’ CSV files onto the laptop.")
step("Put every CSV into the folder named “csvs” inside this "
     "tool’s folder. The filenames don’t matter.")
note("It’s safe to add the same file twice — duplicate readings "
     "are detected and ignored automatically.")

# --- Running ------------------------------------------------------------
doc.add_heading("5. Run the dashboard", level=1)
doc.add_paragraph("Pick whichever matches your laptop:")
bullet(" double-click “run_windows.bat”.", "Windows:")
bullet(" double-click “run_mac.command”.", "Mac:")
bullet(" in the terminal, run the command below.", "Any computer:")
code("python home_base.py csvs")
doc.add_paragraph(
    "After a moment, the dashboard (classroom_dashboard.html) opens in your "
    "web browser. Re-run the launcher whenever more files arrive — it "
    "takes seconds, and the checklist updates to show who is still missing.")
note("To make the checklist expect your full class list:  "
     "python home_base.py csvs --expect 25  (expects Group 1..Group 25), or "
     "--roster roster.txt with your own names one per line. --title "
     "“My School” changes the header title.")

# --- Using it -----------------------------------------------------------
doc.add_heading("6. Using the dashboard in class", level=1)
bullet(" start with a single group (“here’s Group 2”), then "
       "switch Show to “All groups” for the reveal; Outdoor/Indoor "
       "focuses the comparison. “Colour by” picks the variable: "
       "· health bands uses the phones’ absolute colours; · "
       "spread re-stretches to today’s range — use it when "
       "everything looks one colour and structure appears within the green.",
       "Map tab:")
bullet(" the smooth surface is an estimate between the points, fading where "
       "nobody sampled. Widening the radius (tight → wide) fills more "
       "area with more guesswork — the same trade-off NASA’s "
       "GISTEMP temperature maps make (250 vs 1200 km smoothing).",
       "Interpolated tab:")
bullet(" a quick visual of PM2.5 hotspots.", "Heatmap tab:")
bullet(" the indoor-vs-outdoor CO₂ bars are usually the story of the "
       "day; the table lets every group find itself.", "Stats tab:")
doc.add_paragraph(
    "Hover over (or tap) any point for all of its readings. The camera icon "
    "in a tab’s top-right toolbar saves that view as a PNG for slides. "
    "To project, put the browser window on the classroom display (F11 = "
    "full screen on Windows).")

# --- Sharing ------------------------------------------------------------
doc.add_heading("7. Sharing with students afterwards", level=1)
doc.add_paragraph(
    "classroom_dashboard.html is one self-contained file. Email it, post it "
    "to the course page, or put it in a shared folder — it opens in any "
    "browser; internet is needed only for the map background.")
doc.add_paragraph(
    "Privacy: the CSVs contain GPS tracks and the group names students "
    "typed. Keep the csvs folder off shared drives and public sites.")

# --- Troubleshooting ----------------------------------------------------
doc.add_heading("8. Troubleshooting", level=1)
tbl = doc.add_table(rows=1, cols=2)
try:
    tbl.style = "Light Grid Accent 1"
except Exception:
    tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
hdr[0].paragraphs[0].add_run("If you see…").bold = True
hdr[1].paragraphs[0].add_run("Do this").bold = True
rows = [
    ("‘python is not recognized’ (Windows)",
     "Use ‘py’ instead of ‘python’, use the Anaconda Prompt "
     "if you have Anaconda, or reinstall Python and tick ‘Add Python "
     "to PATH’."),
    ("‘Python was not found; run without arguments to install from "
     "the Microsoft Store…’",
     "Windows’ decoy python — no real Python is on your PATH. With "
     "Anaconda: open “Anaconda Prompt” instead, or run by full "
     "path:  C:\\ProgramData\\anaconda3\\python.exe home_base.py csvs  "
     "(Anaconda may also live in %LOCALAPPDATA%\\anaconda3 or "
     "%USERPROFILE%\\anaconda3). Without Anaconda: install Python, "
     "step 3.1."),
    ("‘command not found: python’ (Mac)",
     "Use ‘python3’ (and ‘pip3’)."),
    ("‘No module named plotly/pandas/PIL’",
     "You skipped step 3.2 — run the pip install command in this "
     "tool’s folder."),
    ("An error mentioning ‘Scattermap’",
     "Your Plotly is too old — run:  pip install -U "
     "“plotly>=5.24,<7”."),
    ("Mac: ‘cannot be opened… unidentified developer’",
     "Right-click run_mac.command and choose Open (only needed the first "
     "time)."),
    ("‘No CSVs found’",
     "Put the groups’ CSV files into the ‘csvs’ folder, then "
     "run again."),
    ("⚠ ‘file skipped’ in the dashboard header",
     "That file wasn’t an app export — a raw Temtop download has "
     "no GPS. Ask the group to re-send from the app’s Data screen."),
    ("The map background is blank or grey",
     "No internet — points and stats still work; connect to wifi and "
     "reload for the basemap."),
    ("A group’s points sit far off campus",
     "Their phone gave a coarse (cell-tower) location. Have students enable "
     "Precise Location — and discuss it; real networks screen for "
     "exactly this."),
    ("A group sent data but their chip is grey",
     "The name they typed in the app doesn’t match your roster "
     "spelling — check the green chips for the name they actually "
     "used."),
]
for a, b in rows:
    c = tbl.add_row().cells
    c[0].paragraphs[0].add_run(a)
    c[1].paragraphs[0].add_run(b)
doc.add_paragraph()

# --- Maintainer note ----------------------------------------------------
doc.add_heading("9. For whoever maintains the code", level=1)
bullet(" home_base.py reads the CSVs and writes the dashboard; "
       "make_sample_data.py creates test data; the run_* launchers call "
       "home_base.py. build_map.py is the previous four-file version, kept "
       "unchanged for reference.", "Files:")
bullet(" the colour thresholds live in the ‘VARS’ list near the "
       "top of home_base.py and mirror the app’s "
       "lib/models/map_variable.dart. If a band changes in the app, update "
       "it here too so the projected map keeps matching the phones.",
       "Colours:")
bullet(" GETTING_STARTED.md in the same folder is the on-screen version of "
       "this guide (keep the two in sync); README.md is the short "
       "reference.", "More detail:")

doc.save(OUT)
print("wrote", OUT)
