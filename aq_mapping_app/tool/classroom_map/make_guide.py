"""Generate the Word guide (Classroom_Map_Guide.docx) for TAs / colleagues.

Run:  python3 make_guide.py
Edit the content here and re-run to regenerate the document.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "Classroom_Map_Guide.docx")

doc = Document()

# Base font
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def shade(paragraph, fill="F3F3F3"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def code(text):
    """A shaded monospace command block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    shade(p)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def step(text):
    return doc.add_paragraph(text, style="List Number")


def note(text):
    p = doc.add_paragraph()
    r = p.add_run("Tip:  ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x77, 0x66)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(8)
    return p


# --- Title --------------------------------------------------------------
t = doc.add_heading("AQ Mapper — Classroom Map Guide", level=0)
sub = doc.add_paragraph()
sr = sub.add_run("How to set up and run the air-quality class map. "
                 "Written for a TA or colleague — no programming needed.")
sr.italic = True
sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()

# --- What this tool does ------------------------------------------------
doc.add_heading("1. What this tool does", level=1)
doc.add_paragraph(
    "During the lab, student groups collect air-quality readings on their "
    "phones using the AQ Mapper app and export the data as CSV files. This "
    "tool takes all of those CSV files, combines them, and produces an "
    "interactive map you can project in the classroom for the debrief.")
doc.add_paragraph("On the map you can:")
bullet(" show one group at a time (“here’s Group 2”) or everyone "
       "together, or split by indoor vs. outdoor.", "Filter the data:")
bullet(" PM2.5, PM10, CO₂, HCHO, temperature, or humidity — the "
       "colours match the phone app exactly.", "Colour the points by:")
bullet(" hover over any point to see all of its readings.", "Inspect points:")
bullet(" a second file shows a PM2.5 “heat” map of pollution "
       "hotspots.", "See a heatmap:")
doc.add_paragraph(
    "The output is an ordinary HTML file that opens in any web browser. You "
    "can project it, and you can email it to students so they can explore the "
    "class data on their own afterward.")

# --- What you'll need ---------------------------------------------------
doc.add_heading("2. What you’ll need", level=1)
bullet(" Windows or Mac, either is fine.", "A laptop:")
bullet(" the folder containing build_map.py and the other files (the same "
       "folder this guide came in).", "This tool:")
bullet(" a free program that runs the tool (one-time install, below).",
       "Python 3:")
bullet(" classroom wifi. Only the map’s background needs internet; "
       "everything else works offline.", "Internet:")

# --- One-time setup -----------------------------------------------------
doc.add_heading("3. One-time setup (once per computer)", level=1)

doc.add_heading("3.1  Install Python", level=2)
doc.add_paragraph("Windows:")
bullet(" Open the Microsoft Store, search for “Python 3”, and click "
       "Get/Install. (Or download from python.org and, on the first install "
       "screen, tick “Add Python to PATH”.)")
doc.add_paragraph("Mac:")
bullet(" Python 3 is usually already installed. If not, download it from "
       "python.org and run the installer.")
doc.add_paragraph("To confirm it worked, open a terminal "
                  "(Windows: “Command Prompt”; Mac: “Terminal”) "
                  "and type:")
code("python --version")
note("If Windows says ‘python is not recognized’, try ‘py "
     "--version’. On Mac, use ‘python3 --version’. Use that same "
     "word (python / py / python3) everywhere below.")

doc.add_heading("3.2  Install the two libraries the tool uses", level=2)
doc.add_paragraph(
    "In the terminal, move into this tool’s folder and install its "
    "requirements. To move into the folder, type ‘cd ’ (with a "
    "space) and then drag the folder onto the terminal window and press "
    "Enter. Then run:")
code("pip install -r requirements.txt")
note("If ‘pip’ isn’t found, try ‘pip3 install -r "
     "requirements.txt’ or ‘python -m pip install -r "
     "requirements.txt’. You only ever do this once per computer.")

# --- Each lab: getting the data in -------------------------------------
doc.add_heading("4. Each lab: getting the data in", level=1)
step("Have each group export their data: in the AQ Mapper app, open the Data "
     "screen and tap “Export & Share CSV”. They can AirDrop, email, "
     "or otherwise send you the file.")
step("Collect all the groups’ CSV files onto the laptop.")
step("Put every CSV into the folder named “csvs” inside this tool’s "
     "folder. The filenames don’t matter.")
note("It’s safe to add the same file twice — duplicate readings are "
     "detected and ignored automatically.")

# --- Running the map ----------------------------------------------------
doc.add_heading("5. Run the map", level=1)
doc.add_paragraph("Pick whichever matches your laptop:")
bullet(" double-click “run_windows.bat”.", "Windows:")
bullet(" double-click “run_mac.command”. The first time only, "
       "right-click it and choose Open, then click Open in the dialog (this "
       "tells macOS you trust it).", "Mac:")
bullet(" in the terminal, run the command below.", "Any computer:")
code("python build_map.py csvs")
doc.add_paragraph(
    "After a moment, the map (“classroom_map.html”) opens in your web "
    "browser. A heatmap (“classroom_heatmap.html”) is also created in "
    "the same folder — open it the same way (double-click).")

# --- Using the map ------------------------------------------------------
doc.add_heading("6. Using the map in class", level=1)
doc.add_paragraph("Two dropdown menus sit at the top-left of the map:")
bullet(" choose “All groups”, “Outdoor”, “Indoor”, or "
       "a single group. Pick a group to focus the discussion on their route, "
       "then switch to “All groups” to compare the whole class.",
       "Show:")
bullet(" switch the variable the colours represent (PM2.5, CO₂, "
       "temperature, etc.). Green is good / low; red is high. The colour "
       "thresholds are the same ones shown in the phone app.", "Colour by:")
doc.add_paragraph(
    "Hover over (or tap) any point to see all of that reading’s values. "
    "The colour scale on the right shows the thresholds for the selected "
    "variable. To project, just put the browser window on the classroom "
    "display and, if helpful, press the browser’s full-screen key (F11 on "
    "Windows).")

# --- Sharing ------------------------------------------------------------
doc.add_heading("7. Sharing the map with students", level=1)
doc.add_paragraph(
    "The two HTML files are self-contained. You can email them, post them to "
    "the course page, or put them in a shared folder — students just open "
    "them in any browser. They’ll need internet the first time so the map "
    "background can load.")

# --- Sample data --------------------------------------------------------
doc.add_heading("8. Want to try it first, without real data?", level=1)
doc.add_paragraph(
    "The tool can generate a fake class of five groups around UTSC so you can "
    "practise before the lab. In the terminal, run:")
code("python make_sample_data.py\npython build_map.py sample_csvs")
doc.add_paragraph("This builds the same map from invented data. Real data "
                  "always goes in the “csvs” folder instead.")

# --- Troubleshooting ----------------------------------------------------
doc.add_heading("9. Troubleshooting", level=1)
tbl = doc.add_table(rows=1, cols=2)
try:
    tbl.style = "Light Grid Accent 1"
except Exception:
    tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
hdr[0].paragraphs[0].add_run("If you see…").bold = True
hdr[1].paragraphs[0].add_run("Do this").bold = True
rows = [
    ("‘python is not recognized’ (Windows)",
     "Use ‘py’ instead of ‘python’, or reinstall Python and "
     "tick ‘Add Python to PATH’."),
    ("‘command not found: python’ (Mac)",
     "Use ‘python3’ instead of ‘python’."),
    ("‘No module named plotly/pandas’",
     "You skipped step 3.2 — run the pip install command in this tool’s "
     "folder."),
    ("Mac: ‘cannot be opened… unidentified developer’",
     "Right-click run_mac.command and choose Open (only needed the first "
     "time)."),
    ("‘No CSVs found’",
     "Put the groups’ CSV files into the ‘csvs’ folder, then run "
     "again."),
    ("The map background is blank or grey",
     "No internet — connect to wifi and reopen the map. The points still "
     "work offline; only the background needs a connection."),
    ("Some points are missing",
     "A file must be an AQ Mapper export (it contains GPS coordinates). A raw "
     "sensor file straight off the Temtop has no location and is skipped."),
]
for a, b in rows:
    c = tbl.add_row().cells
    c[0].paragraphs[0].add_run(a)
    c[1].paragraphs[0].add_run(b)
doc.add_paragraph()

# --- Maintainer note ----------------------------------------------------
doc.add_heading("10. For whoever maintains the code", level=1)
bullet(" build_map.py reads the CSVs and writes the maps; make_sample_data.py "
       "creates test data; the run_* launchers just call build_map.py.",
       "Files:")
bullet(" the colour thresholds live in the ‘VARS’ list near the top "
       "of build_map.py and mirror the app’s lib/models/map_variable.dart. "
       "If a band changes in the app, update it here too so the projected map "
       "keeps matching the phones.", "Colours:")
bullet(" a short version of this guide is in README.md in the same folder.",
       "More detail:")

doc.save(OUT)
print("wrote", OUT)
