"""Generate the student handout: a QR code for the app plus a one-page card.

Run:  python3 make_handout.py
Outputs aq_mapper_qr.png and AQ_Mapper_Student_Handout.docx in this folder.
Edit APP_URL / wording here and re-run.
"""
import os
import qrcode
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
APP_URL = "https://danweaver-ca.github.io/aq-mapper/"
QR_PNG = os.path.join(HERE, "aq_mapper_qr.png")
DOCX = os.path.join(HERE, "AQ_Mapper_Student_Handout.docx")
TEAL = "#009688"

# --- QR code -------------------------------------------------------------
qr = qrcode.QRCode(error_correction=qrcode.constants.ERROR_CORRECT_M,
                   box_size=12, border=2)
qr.add_data(APP_URL)
qr.make(fit=True)
qr.make_image(fill_color=TEAL, back_color="white").save(QR_PNG)
print("wrote", QR_PNG)

# --- one-page handout ----------------------------------------------------
doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


h = doc.add_heading("AQ Mapper — Air Quality Lab", level=0)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Log air-quality readings on your phone — no app to install.")
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(QR_PNG, width=Inches(2.4))
link = doc.add_paragraph()
link.alignment = WD_ALIGN_PARAGRAPH.CENTER
lr = link.add_run(APP_URL)
lr.bold = True
lr.font.size = Pt(13)
doc.add_paragraph()

doc.add_heading("Getting started", level=1)
for i, text in enumerate([
    "Scan the QR code with your phone camera (or type the link above). The app "
    "opens right in your web browser — nothing to download.",
    "Recommended — keep it handy: iPhone → tap the Share button → "
    "“Add to Home Screen.”  Android → tap the menu (⋮) → “Add to Home "
    "Screen” / “Install.” Then open it from the new icon.",
    "Enter your Group name and device ID when prompted.",
    "Tap “Add Measurement” to log a reading. Allow location access when "
    "asked so each reading is placed on the map.",
], start=1):
    doc.add_paragraph(text, style="List Number")

# Warning box
doc.add_paragraph()
warn = doc.add_paragraph()
shade(warn, "FFF3CD")
warn.paragraph_format.space_before = Pt(6)
warn.paragraph_format.space_after = Pt(6)
wr = warn.add_run("⚠  Export before you leave!  ")
wr.bold = True
wr.font.color.rgb = RGBColor(0x8A, 0x61, 0x00)
warn.add_run(
    "Your readings are saved only on your phone. Before you finish, open the "
    "Data screen, tap “Export & Share CSV,” and send the file to your "
    "instructor. Closing or clearing your browser can erase the data.")

tip = doc.add_paragraph()
tr = tip.add_run("Tip: the campus map works even without wifi, but it needs a "
                 "connection the first time you open the app.")
tr.italic = True
tr.font.size = Pt(10)
tr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.save(DOCX)
print("wrote", DOCX)
